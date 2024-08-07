import os
from docx import Document

def merge_txt_files_to_word(directory):
    # Ensure the directory exists
    if not os.path.isdir(directory):
        print("The specified directory does not exist.")
        return
    
    # List all .txt files in the directory
    txt_files = [f for f in os.listdir(directory) if f.endswith('.txt')]
    
    if not txt_files:
        print("No .txt files found in the directory.")
        return
    
    # Create a Document object
    doc = Document()
    
    # Iterate over all .txt files and add their content to the Word document
    for txt_file in txt_files:
        file_path = os.path.join(directory, txt_file)
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            # Add the filename as a heading (optional)
            doc.add_heading(txt_file, level=1)
            # Add the content
            doc.add_paragraph(content)
    
    # Define the output Word file path
    output_path = os.path.join(directory, 'merged_output.docx')
    
    # Save the Word document
    doc.save(output_path)
    print(f"All .txt files have been merged into {output_path}")

# Example usage
directory_path = input("Enter the path of the directory containing .txt files: ")
merge_txt_files_to_word(directory_path)
